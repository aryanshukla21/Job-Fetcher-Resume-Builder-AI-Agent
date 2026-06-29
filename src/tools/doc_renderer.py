"""
Document Generator and Renderer.
Manipulates document trees at the run level to enforce visual style persistence.
"""
from __future__ import annotations

import os
from pathlib import Path
import docx

from src.schemas.models import ResumeDraftContent

def compile_tailored_resume(template_path: str, output_path: str, modifications: ResumeDraftContent) -> str:
    """
    Reads a master .docx template, mutates text elements safely without corrupting font metadata, 
    and writes out the tailored production version.
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Master resume template not found at: {template_path}")

    doc = docx.Document(template_path)

    # 1. Dynamic Objective/Summary Text Re-alignment
    # Searches for a structural placeholder string safely embedded inside the template run trees
    _replace_placeholder_in_place(doc, "{{PROFESSIONAL_SUMMARY}}", modifications.summary)

    # 2. Targeted Bullets Tailoring
    # Matches and mutates individual project/experience lines without dropping layout characteristics
    _tailor_experience_bullets(doc, modifications)

    # Ensure targeted output directory trees exist before executing write operations
    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    
    doc.save(str(output_file))
    return str(output_file.resolve())

def _replace_placeholder_in_place(doc: docx.Document, placeholder: str, replacement_text: str) -> None:
    """Safely updates text segments at the run level to maintain original styling rules."""
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, replacement_text)
                    
    # Also parse table cells for placeholder targets
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, replacement_text)

def _tailor_experience_bullets(doc: docx.Document, modifications: ResumeDraftContent) -> None:
    """
    Finds bullet placeholder hooks and maps updated descriptions onto the document structure.
    """
    # Maps internal identifiers to specific update pools provided by the agent execution state
    # Assumes your template file uses tracking tokens like {{EXP_BULLET_1}}, {{PROJECT_BULLET_1}}
    bullet_index = 1
    
    # Process structured job modifications
    for job in modifications.updated_experiences:
        for bullet in job.highlights:
            hook = f"{{{{EXP_BULLET_{bullet_index}}}}}"
            _replace_placeholder_in_place(doc, hook, bullet)
            bullet_index += 1

    # Process structured project modifications
    proj_index = 1
    for proj in modifications.updated_projects:
        for bullet in proj.highlights:
            hook = f"{{{{PROJ_BULLET_{proj_index}}}}}"
            _replace_placeholder_in_place(doc, hook, bullet)
            proj_index += 1